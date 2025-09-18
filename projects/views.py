from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from .forms import LLMProjectForm
from .models import LLMProject, ProjectDocument
from django.db.models import Q
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_POST
from django.template.loader import render_to_string
import re
import requests
from requests.exceptions import SSLError
from bs4 import BeautifulSoup
import os
import io
import json
import zipfile
from django.utils import timezone
import google.generativeai as genai
from .llm import call_gemini, call_gemini_image
from .retriever import search_similar_docs



@login_required
def project_test(request, pk):
    """
    popup 頁面：顯示問題輸入框 + 範例提示詞清單
    """
    obj = get_object_or_404(LLMProject, pk=pk)
    examples_raw = getattr(obj, "example_prompts", "") or ""
    # 把例子逐行拆成 list，忽略空行
    examples = [line.strip() for line in examples_raw.splitlines() if line.strip()]
    return render(request, "projects/project_test.html", {"obj": obj, "examples": examples})

@require_POST
@login_required
def project_test_api(request, pk):
    """
    AJAX endpoint：接收 question，回傳 JSON
    先用向量檢索找出相關資料，再交給 Gemini 回答
    """
    obj = get_object_or_404(LLMProject, pk=pk)
    question = (request.POST.get("question") or "").strip()
    if not question:
        return JsonResponse({"error": "請輸入問題"}, status=400)

    # 1) 檢索最相近的知識片段
    hits = search_similar_docs(project_id=obj.id, question=question, top_k=3)

    # 2) 把片段串成 context（可加上來源標號，方便你前端顯示）
    if hits:
        context_text = "\n\n".join([f"[{i+1}] {h['text']}" for i, h in enumerate(hits)])
    else:
        context_text = ""

    # 3) 角色指令（維持你的設定，也可加上口吻/格式要求）
    # role_prompt = "你是一位專業的人事問答助理。請優先依據已知資料回答，必要時再補充一般常識。"
    role_prompt =  obj.role_prompt+' '+obj.response_template 

    # 4) 呼叫 LLM
    answer = call_gemini(role_prompt, question, context=context_text)

    # 5) 回傳答案 +（可選）回傳命中的片段與分數，方便前端顯示/除錯
    return JsonResponse({
        'answer': answer,
        'sources': hits,  # ← 前端可顯示來源與相似度
    })


@require_POST
@login_required
def project_generate_image_api(request, pk):
    """
    依據前端送來的 img_prompt 產生一張圖片，回傳 base64（data URL）。
    """
    _ = get_object_or_404(LLMProject, pk=pk)
    img_prompt = (request.POST.get('img_prompt') or '').strip()
    if not img_prompt:
        return JsonResponse({'error': '缺少 img_prompt'}, status=400)

    try:
        img_bytes, mime_type = call_gemini_image(img_prompt)

        import base64
        b64 = base64.b64encode(img_bytes).decode('ascii')
        return JsonResponse({'image_base64': b64, 'content_type': mime_type})

        # 前端可直接使用 data:image/png;base64,<b64>
        #return JsonResponse({'image_base64': b64, 'content_type': 'image/png'})
    except Exception as e:
        return JsonResponse({'error': f'圖片生成失敗：{e}'}, status=500)






@login_required
def project_new(request):
    if request.method == "POST":
        form = LLMProjectForm(request.POST)
        if form.is_valid():
            obj = form.save()
            return redirect("project_detail", pk=obj.pk)
    else:
        form = LLMProjectForm()
    return render(request, "projects/project_form.html", {"form": form})

@login_required
def project_detail(request, pk):
    obj = get_object_or_404(LLMProject, pk=pk)
    return render(request, "projects/project_detail.html", {"obj": obj})

# ① 先搜尋頁（/project/edit?q=關鍵字）
@login_required
def project_edit_search(request):
    q = (request.GET.get("q") or "").strip()
    if q:
        results = (
            LLMProject.objects
            .filter(Q(project_code__icontains=q) | Q(name__icontains=q))
            .order_by("project_code", "name")
        )
    else:
        # ✅ 如果沒輸入，顯示全部
        results = LLMProject.objects.all().order_by("project_code", "name")

    return render(
        request,
        "projects/project_edit_search.html",
        {"q": q, "results": results},
    )
# ① 先搜尋頁（/project/edit?q=關鍵字）
@login_required
def project_publish(request):
    q = (request.GET.get("q") or "").strip()
    if q:
        results = (
            LLMProject.objects
            .filter(Q(project_code__icontains=q) | Q(name__icontains=q))
            .order_by("project_code", "name")
        )
    else:
        # ✅ 如果沒輸入，顯示全部
        results = LLMProject.objects.all().order_by("project_code", "name")

    return render(
        request,
        "projects/project_publish.html",
        {"q": q, "results": results},
    )

# ② 編輯頁（/project/edit/<pk>）
@login_required
def project_edit(request, pk):
    obj = get_object_or_404(LLMProject, pk=pk)
    if request.method == "POST":
        form = LLMProjectForm(request.POST, instance=obj)
        if form.is_valid():
            form.save()
            return redirect("project_detail", pk=obj.pk)
    else:
        form = LLMProjectForm(instance=obj)
    # 重用和 new 相同的表單模板，只是換標題
    return render(
        request,
        "projects/project_form.html",
        {"form": form, "title": f"編輯專案：{obj.project_code or obj.name}", "is_edit": True, "obj": obj},
    )













# 第一步：檔案內容解析 helper
def extract_text_from_file(file_obj, filename):
    """根據副檔名嘗試抽出文字（簡單實作）。
    需要的第三方套件：PyPDF2, python-docx, pandas, openpyxl
    若缺少套件，會 raise ImportError，請按照錯誤安裝。
    """
    ext = os.path.splitext(filename)[1].lower()
    file_obj.seek(0)

    if ext == '.pdf':
        try:
            from PyPDF2 import PdfReader
        except Exception as e:
            raise ImportError('請安裝 PyPDF2 (pip install PyPDF2)')
        reader = PdfReader(file_obj)
        texts = []
        for page in reader.pages:
            texts.append(page.extract_text() or '')
        return '\n'.join(texts)

    if ext in ('.docx', '.doc'):
        try:
            import docx
        except Exception:
            raise ImportError('請安裝 python-docx (pip install python-docx)')
        # docx (python-docx) 只支援 .docx
        if ext == '.docx':
            doc = docx.Document(file_obj)
            paragraphs = [p.text for p in doc.paragraphs]
            return '\n'.join(paragraphs)
        else:
            # .doc 的處理可用 textract 或 antiword — 這裡回傳空字串並提醒安裝對應工具
            return ''

    if ext in ('.xls', '.xlsx'):
        try:
            import pandas as pd
        except Exception:
            raise ImportError('請安裝 pandas 與 openpyxl/xlrd (pip install pandas openpyxl xlrd)')
        # 使用 pandas 讀取所有 sheet 並串接文字
        try:
            # pandas can read file-like for read_excel
            xls = pd.read_excel(file_obj, sheet_name=None)
            parts = []
            for sheet_name, df in xls.items():
                parts.append(f'-- Sheet: {sheet_name} --')
                parts.append(df.astype(str).apply(lambda row: '\t'.join(row.values), axis=1).str.cat(sep='\n'))
            return '\n'.join(parts)
        except Exception:
            file_obj.seek(0)
            return ''

    if ext == '.csv':
        try:
            import pandas as pd
        except Exception:
            raise ImportError('請安裝 pandas (pip install pandas)')
        file_obj.seek(0)
        df = pd.read_csv(file_obj)
        return df.astype(str).apply(lambda row: '\t'.join(row.values), axis=1).str.cat(sep='\n')

    # 預設當成純文字
    try:
        file_obj.seek(0)
        raw = file_obj.read()
        # file_obj.read() 可能回傳 bytes
        if isinstance(raw, bytes):
            return raw.decode('utf-8', errors='ignore')
        return str(raw)
    except Exception:
        return ''


# 第二步：向量化 helper（這裡放一個 placeholder）
def compute_embedding(text):
    try:
        # 1. 設定 API 金鑰
        # 如果你使用環境變數，將 settings.GOOGLE_API_KEY 改成 os.getenv("GOOGLE_API_KEY")
        genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
        
        # 2. 選擇向量化模型
        # 'models/text-embedding-004' 是 Google 推薦的向量化模型
        model = 'models/text-embedding-004'
        
        # 3. 呼叫向量化模型
        response = genai.embed_content(
            model=model,
            content=text,
            task_type="retrieval_document" # 這裡設定任務類型，可以提高準確性
        )
        
        # 4. 從回應中取出向量值（通常是一個列表）
        # 'embedding' 欄位包含了我們需要的向量
        return response['embedding']

    except Exception as e:
        print(f"向量化失敗: {e}")
        # 如果發生錯誤，返回一個空列表或你設定的預設值
        return []

# 第三步： import page（GET 顯示上傳表單與已匯入列表；POST 處理上傳）
@login_required
def project_import(request, pk):
    project = get_object_or_404(LLMProject, pk=pk)

    if request.method == 'POST':
        # 可選擇上傳檔案、手動輸入，或（可選）伺服器路徑
        uploaded = request.FILES.get('file')
        path_text = request.POST.get('path')
        action = request.POST.get('action')
        manual_title = (request.POST.get('manual_title') or '').strip()
        manual_content = (request.POST.get('manual_content') or '').strip()
        
        # Debug: 檢查是否有檔案上傳
        print(f"DEBUG: uploaded = {uploaded}")
        print(f"DEBUG: path_text = {path_text}")
        print(f"DEBUG: request.FILES = {request.FILES}")
        print(f"DEBUG: request.POST = {request.POST}")

        # 優先處理手動輸入
        if action == 'manual_save' and manual_content:
            filename = manual_title or (f"manual_{timezone.now().strftime('%Y%m%d_%H%M%S')}.txt")
            doc = ProjectDocument(project=project, imported_by=request.user)
            doc.filename = filename
            doc.content = manual_content
            try:
                doc.embedding = compute_embedding(manual_content)
            except Exception:
                doc.embedding = None
            doc.save()
            return redirect('project_import', pk=project.pk)

        if not uploaded and not path_text and not manual_content:
            return render(request, 'projects/project_import.html', {
                'project': project, 'error': '請上傳檔案',
                'documents': project.documents.all().order_by('-imported_at')
            })

        # 儲存 model
        doc = ProjectDocument(project=project, imported_by=request.user)
        if uploaded:
            doc.filename = uploaded.name
            doc.uploaded_file = uploaded
            # 立即解析內容（檔案已存在 uploaded_file），但使用 Django 上傳的 file-like
            file_like = uploaded.file if hasattr(uploaded, 'file') else uploaded
            try:
                content = extract_text_from_file(file_like, uploaded.name)
            except ImportError as e:
                # 缺套件
                return render(request, 'projects/project_import.html', {
                    'project': project, 'error': str(e),
                    'documents': project.documents.all().order_by('-imported_at')
                })
        else:
            # 若輸入伺服器路徑（例：/data/files/doc1.pdf）
            doc.filename = os.path.basename(path_text)
            try:
                with open(path_text, 'rb') as f:
                    content = extract_text_from_file(f, doc.filename)
            except Exception as e:
                return render(request, 'projects/project_import.html', {
                    'project': project, 'error': f'讀取路徑失敗：{e}',
                    'documents': project.documents.all().order_by('-imported_at')
                })

        # 儲存文字與 embedding
        doc.content = content
        try:
            doc.embedding = compute_embedding(content)
        except Exception:
            doc.embedding = None
        doc.save()

        return redirect('project_import', pk=project.pk)

    # GET 顯示
    documents = project.documents.all().order_by('-imported_at')
    return render(request, 'projects/project_import.html', {
        'project': project,
        'documents': documents,
    })


# 第四步： 檢視已匯入資料
@login_required
def project_import_detail(request, pk, doc_pk):
    project = get_object_or_404(LLMProject, pk=pk)
    doc = get_object_or_404(ProjectDocument, pk=doc_pk, project=project)
    if request.method == 'POST':
        new_filename = (request.POST.get('filename') or doc.filename).strip()
        new_content = (request.POST.get('content') or '').strip()

        doc.filename = new_filename or doc.filename
        doc.content = new_content
        try:
            doc.embedding = compute_embedding(new_content)
        except Exception:
            doc.embedding = None
        doc.save()

        return redirect('project_import_detail', pk=project.pk, doc_pk=doc.pk)

    return render(request, 'projects/project_import_detail.html', {
        'project': project,
        'doc': doc,
    })


# 第五步： 刪除
@login_required
def project_import_delete(request, pk, doc_pk):
    project = get_object_or_404(LLMProject, pk=pk)
    doc = get_object_or_404(ProjectDocument, pk=doc_pk, project=project)
    if request.method == 'POST':
        doc.delete()
        return redirect('project_import', pk=project.pk)
    return render(request, 'projects/project_import_delete_confirm.html', {'project': project, 'doc': doc})


# 匯出：將此專案的 ProjectDocument 以 SQL 產生並打包 ZIP 下載
@login_required
def project_export_sql(request, pk):
    project = get_object_or_404(LLMProject, pk=pk)

    # 準備 SQL 內容（SQLite/MySQL 皆可接受的基本 INSERT 語法）。
    # 資料表名依 Django 預設為 app_label_modelname → projects_projectdocument
    table_name = 'projects_projectdocument'

    def sql_value(v):
        if v is None:
            return 'NULL'
        if isinstance(v, (int, float)):
            return str(v)
        # 其他型別一律當成字串，做基本跳脫：單引號→兩個單引號，換行→\n
        s = v
        if not isinstance(s, str):
            s = str(s)
        s = s.replace("'", "''").replace("\\", "\\\\").replace("\n", r"\n").replace("\r", r"\r")
        return f"'{s}'"

    rows = (
        ProjectDocument.objects
        .filter(project=project)
        .order_by('id')
        .values(
            'id', 'project_id', 'filename', 'uploaded_file', 'content', 'embedding', 'imported_by_id', 'imported_at'
        )
    )

    lines = []
    lines.append('-- Exported ProjectDocument rows for project id=%d (%s)' % (project.id, project.project_code or project.name))
    lines.append('-- Generated at %s' % timezone.now().strftime('%Y-%m-%d %H:%M:%S'))
    lines.append('BEGIN;')

    # 可選：先刪除同專案舊資料，以避免重複（如不需要可移除下一行）
    lines.append(f"DELETE FROM {table_name} WHERE project_id = {project.id};")

    if rows:
        columns = ['id', 'project_id', 'filename', 'uploaded_file', 'content', 'embedding', 'imported_by_id', 'imported_at']
        col_list = ', '.join(columns)
        for r in rows:
            values = []
            for c in columns:
                val = r[c]
                if c == 'embedding' and val is not None:
                    # 序列化為 JSON 字串
                    val = json.dumps(val, ensure_ascii=False)
                if c == 'imported_at' and val is not None:
                    # 轉為 ISO 格式字串
                    val = val.isoformat(sep=' ', timespec='seconds')
                values.append(sql_value(val))
            values_sql = ', '.join(values)
            lines.append(f"INSERT INTO {table_name} ({col_list}) VALUES ({values_sql});")

    lines.append('COMMIT;')

    sql_bytes = '\n'.join(lines).encode('utf-8')

    # 以 ZIP 打包
    memfile = io.BytesIO()
    with zipfile.ZipFile(memfile, mode='w', compression=zipfile.ZIP_DEFLATED) as zf:
        filename_part = (project.project_code or project.name or f'project_{project.id}').replace(' ', '_')
        sql_name = f"project_{filename_part}_documents.sql"
        zf.writestr(sql_name, sql_bytes)

    memfile.seek(0)
    resp = HttpResponse(memfile.read(), content_type='application/zip')
    ts = timezone.now().strftime('%Y%m%d_%H%M%S')
    download_name = f"project_{project.id}_documents_{ts}.zip"
    resp['Content-Disposition'] = f'attachment; filename="{download_name}"'
    return resp


# 匯出：將此專案的 LLMProject 一筆資料以 SQL 產生並打包 ZIP 下載
@login_required
def project_export_project_sql(request, pk):
    project = get_object_or_404(LLMProject, pk=pk)

    table_name = 'projects_llmproject'

    def sql_value(v):
        if v is None:
            return 'NULL'
        if isinstance(v, (int, float)):
            return str(v)
        s = v
        if not isinstance(s, str):
            s = str(s)
        s = s.replace("'", "''").replace("\\", "\\\\").replace("\n", r"\n").replace("\r", r"\r")
        return f"'{s}'"

    columns = [
        'id', 'project_code', 'name', 'description', 'llm_model', 'embedding_engine',
        'role_prompt', 'response_template', 'example_prompts', 'created_at'
    ]

    data = {
        'id': project.id,
        'project_code': project.project_code,
        'name': project.name,
        'description': project.description,
        'llm_model': project.llm_model,
        'embedding_engine': project.embedding_engine,
        'role_prompt': project.role_prompt,
        'response_template': project.response_template,
        'example_prompts': project.example_prompts,
        'created_at': project.created_at.isoformat(sep=' ', timespec='seconds') if project.created_at else None,
    }

    lines = []
    lines.append('-- Exported LLMProject row id=%d (%s)' % (project.id, project.project_code or project.name))
    lines.append('-- Generated at %s' % timezone.now().strftime('%Y-%m-%d %H:%M:%S'))
    lines.append('BEGIN;')
    # 可選：先刪除同 id 的舊資料
    lines.append(f"DELETE FROM {table_name} WHERE id = {project.id};")

    col_list = ', '.join(columns)
    values_sql = ', '.join(sql_value(data[c]) for c in columns)
    lines.append(f"INSERT INTO {table_name} ({col_list}) VALUES ({values_sql});")
    lines.append('COMMIT;')

    sql_bytes = '\n'.join(lines).encode('utf-8')

    memfile = io.BytesIO()
    with zipfile.ZipFile(memfile, mode='w', compression=zipfile.ZIP_DEFLATED) as zf:
        filename_part = (project.project_code or project.name or f'project_{project.id}').replace(' ', '_')
        sql_name = f"project_{filename_part}.sql"
        zf.writestr(sql_name, sql_bytes)

    memfile.seek(0)
    resp = HttpResponse(memfile.read(), content_type='application/zip')
    ts = timezone.now().strftime('%Y%m%d_%H%M%S')
    download_name = f"project_{project.id}_{ts}.zip"
    resp['Content-Disposition'] = f'attachment; filename="{download_name}"'
    return resp


# 匯出：將 project_test.html 渲染為靜態 HTML，打包 ZIP 下載
@login_required
def project_export_example_html(request, pk):
    obj = get_object_or_404(LLMProject, pk=pk)
    examples_raw = getattr(obj, "example_prompts", "") or ""
    examples = [line.strip() for line in examples_raw.splitlines() if line.strip()]

    html = render_to_string('projects/project_test.html', {"obj": obj, "examples": examples, "request": request})

    memfile = io.BytesIO()
    with zipfile.ZipFile(memfile, mode='w', compression=zipfile.ZIP_DEFLATED) as zf:
        filename_part = (obj.project_code or obj.name or f'project_{obj.id}').replace(' ', '_')
        zf.writestr(f"{filename_part}_project_test.html", html.encode('utf-8'))

    memfile.seek(0)
    resp = HttpResponse(memfile.read(), content_type='application/zip')
    ts = timezone.now().strftime('%Y%m%d_%H%M%S')
    download_name = f"project_{obj.id}_example_html_{ts}.zip"
    resp['Content-Disposition'] = f'attachment; filename="{download_name}"'
    return resp


# 爬蟲：輸入網址，擷取文字，轉為 Markdown，向量化後匯入資料庫
@login_required
def project_crawl(request, pk):
    project = get_object_or_404(LLMProject, pk=pk)

    context = {
        'project': project,
        'result': None,
        'error': None,
    }

    if request.method == 'POST':
        url = (request.POST.get('url') or '').strip()
        if not url:
            context['error'] = '請輸入網址'
            return render(request, 'projects/project_crawl.html', context)

        try:
            resp = requests.get(
                url,
                timeout=20,
                headers={'User-Agent': 'Mozilla/5.0 (compatible; CSWBot/1.0; +https://example.com/bot)'},
                verify=True,
            )
            resp.raise_for_status()
        except SSLError as e:
            context['error'] = f'抓取失敗（SSL）：{e}'
            return render(request, 'projects/project_crawl.html', context)
        except Exception as e:
            context['error'] = f'抓取失敗：{e}'
            return render(request, 'projects/project_crawl.html', context)

        # 解析 HTML，抽取主要文字
        html = resp.text
        soup = BeautifulSoup(html, 'html.parser')
        # 移除 script/style
        for tag in soup(['script', 'style', 'noscript']):
            tag.decompose()
        # 嘗試抓主體
        main = soup.find('main') or soup.find('article') or soup
        text_parts = []
        # 標題
        if soup.title and soup.title.string:
            text_parts.append(f"# {soup.title.string.strip()}")
        # 轉為 Markdown-ish（簡化版）
        for h in main.find_all(['h1','h2','h3','h4','h5','h6']):
            prefix = '#' * int(h.name[1])
            text_parts.append(f"\n{prefix} {h.get_text(strip=True)}\n")
        # 段落
        for p in main.find_all('p'):
            content = p.get_text(" ", strip=True)
            if content:
                text_parts.append(content)
        # 清理多重空白
        md_text = '\n\n'.join(text_parts)
        md_text = re.sub(r'\n{3,}', '\n\n', md_text).strip()
        if not md_text:
            md_text = soup.get_text("\n", strip=True)

        # 儲存為 ProjectDocument
        filename = re.sub(r'[^a-zA-Z0-9._-]+', '_', url)[:200] or 'page'
        doc = ProjectDocument(
            project=project,
            filename=f"crawl_{filename}.md",
            imported_by=request.user,
            content=md_text,
        )
        try:
            doc.embedding = compute_embedding(md_text)
        except Exception:
            doc.embedding = None
        doc.save()

        context['result'] = {
            'doc_id': doc.id,
            'filename': doc.filename,
        }
        return render(request, 'projects/project_crawl.html', context)

    return render(request, 'projects/project_crawl.html', context)
